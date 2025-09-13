"""
File Handler Utility
Handles file uploads, validation, text extraction from CVs
"""

import os
import logging
from typing import Optional
import PyPDF2
from docx import Document
from werkzeug.utils import secure_filename
from flask import current_app

logger = logging.getLogger(__name__)

class FileHandler:
    def __init__(self):
        """Initialize file handler"""
        self.upload_folder = 'uploads'
        self.allowed_extensions = {'pdf', 'doc', 'docx'}
        self.max_file_size = 16 * 1024 * 1024  # 16MB
        
        # Create upload directory if it doesn't exist
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def save_file(self, file, filename: str) -> str:
        """
        Save uploaded file to upload directory
        
        Args:
            file: Uploaded file object
            filename: Original filename
            
        Returns:
            Path to saved file
        """
        try:
            # Secure the filename
            secure_name = secure_filename(filename)
            
            # Generate unique filename to avoid conflicts
            import uuid
            unique_id = str(uuid.uuid4())[:8]
            name, ext = os.path.splitext(secure_name)
            unique_filename = f"{name}_{unique_id}{ext}"
            
            # Full file path
            file_path = os.path.join(self.upload_folder, unique_filename)
            
            # Save the file
            file.save(file_path)
            
            logger.info(f"File saved successfully: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from uploaded file
        
        Args:
            file_path: Path to the uploaded file
            
        Returns:
            Extracted text content
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append(text)
            
            # Join all pages
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                raise Exception("No text could be extracted from the PDF")
            
            return full_text
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {str(e)}")
            raise Exception(f"Could not extract text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                raise Exception("No text could be extracted from the document")
            
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise Exception(f"Could not extract text from document: {str(e)}")
    
    def cleanup_file(self, file_path: str) -> bool:
        """
        Clean up uploaded file after processing
        
        Args:
            file_path: Path to file to delete
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File cleaned up: {file_path}")
                return True
            return False
            
        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {str(e)}")
            return False
    
    def validate_file(self, file) -> tuple:
        """
        Validate uploaded file
        
        Args:
            file: Uploaded file object
            
        Returns:
            Tuple of (is_valid: bool, error_message: str)
        """
        if not file:
            return False, "No file provided"
        
        if file.filename == '':
            return False, "No file selected"
        
        if not self.allowed_file(file.filename):
            return False, f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
        
        # Check file size if possible
        if hasattr(file, 'content_length') and file.content_length:
            if file.content_length > self.max_file_size:
                max_size_mb = self.max_file_size / (1024 * 1024)
                return False, f"File too large. Maximum size: {max_size_mb}MB"
        
        return True, ""
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get information about a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        try:
            if not os.path.exists(file_path):
                return {}
            
            file_stat = os.stat(file_path)
            file_info = {
                'size': file_stat.st_size,
                'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
                'modified': file_stat.st_mtime,
                'extension': os.path.splitext(file_path)[1].lower(),
                'filename': os.path.basename(file_path)
            }
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {}
